"""
Directive 78: LOI PDF Automator

Ingests the D67 outreach CRM CSV and LOI markdown templates,
fills blanks from parcel data, outputs ready-to-sign DOCX files.

Usage:
  python3 data_center/forge_lois.py [--all] [--top N]
"""

import argparse
import csv
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = Path(__file__).parent.parent / "output"
CRM_PATH = OUTPUT_DIR / "d67_outreach_crm.csv"
LOI_DIR = OUTPUT_DIR / "d78_generated_lois"

# Price per acre tiers (data center / industrial development)
PRICE_PER_ACRE = {
    'INDUSTRIAL': 3500,
    'COMMERCIAL': 2500,
    'AGRICULTURAL': 1500,
    'RESIDENTIAL': 1000,
}


def load_crm(path):
    """Load CRM CSV, return list of dicts."""
    with open(path, newline='') as f:
        return list(csv.DictReader(f))


def purchase_price(acres, zoning):
    ppa = PRICE_PER_ACRE.get(zoning, 2000)
    return round(acres * ppa, 2)


def md_to_docx(md_text, output_path):
    """Render simplified markdown to DOCX using python-docx."""
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for line in md_text.split('\n'):
        line_stripped = line.strip()

        # Skip HR
        if line_stripped.startswith('---'):
            continue

        # Table rows
        if line_stripped.startswith('|'):
            if line_stripped.startswith('|---'):
                continue
            cells = [c.strip() for c in line_stripped.split('|')[1:-1]]
            # We simplify tables as text lines for DOCX
            doc.add_paragraph(' | '.join(cells), style='Normal')
            continue

        if not line_stripped:
            continue

        # Headings
        if line_stripped.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped[3:])
            run.bold = True
            run.font.size = Pt(14)
            continue
        if line_stripped.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(line_stripped[2:])
            run.bold = True
            run.font.size = Pt(16)
            continue

        # Bold markers
        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', line_stripped)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    doc.save(str(output_path))


def fill_loi_target(target, loi_num):
    """Fill the LOI template for one target and save DOCX."""
    owner = target['owner_name'].strip()
    acres = float(target['total_acres'] if target.get('total_acres') else target.get('primary_acres', '0'))
    zoning = target.get('primary_zoning', 'INDUSTRIAL')
    dist = target.get('dist_345kv_mi', '0')
    lat = target.get('primary_lat', '0')
    lng = target.get('primary_lng', '0')
    oid = target.get('primary_parcel_oid', 'UNKNOWN')
    synergy = target.get('synergy_score', '0')

    price = purchase_price(acres, zoning)
    price_fmt = f"${price:,.0f}"
    ppa = PRICE_PER_ACRE.get(zoning, 2000)
    ppa_fmt = f"${ppa:,}/acre"

    # Determine county from lat/lng (rough heuristics)
    lat_f = float(lat)
    lng_f = float(lng)
    if 32.5 <= lat_f <= 33.2 and -97.5 <= lng_f <= -96.5:
        county = 'Dallas'
        ercot = 'HIGH'
        ercot_zone = 'North (Dallas Load Zone)'
    elif 30.0 <= lat_f <= 30.5 and -98.2 <= lng_f <= -97.5:
        county = 'Travis'
        ercot = 'HIGH'
        ercot_zone = 'South (Austin Load Zone)'
    elif 31.5 <= lat_f <= 32.2 and -99.0 <= lng_f <= -98.5:
        county = 'Brown/Comanche'
        ercot = 'LOW (ALPHA)'
        ercot_zone = 'West (Frontier)'
    else:
        county = 'Texas'
        ercot = 'LOW (ALPHA)'
        ercot_zone = 'ERCOT'

    # Determine use type
    use_map = {
        'INDUSTRIAL': 'Industrial / Data Center Development',
        'COMMERCIAL': 'Commercial Development',
        'AGRICULTURAL': 'Land Conversion / Data Center Development',
        'RESIDENTIAL': 'Land Conversion',
    }
    land_use = use_map.get(zoning, 'Data Center Development')

    template = f"""# Letter of Intent — Non-Binding Offer to Purchase

**Date:** June 2026
**To:** {owner}
**RE:** Approximately {acres:.0f} acres — {land_use}
**Property OID:** {oid}
**County:** {county}, Texas

---

## 1. Offer Summary

[BUYER ENTITY NAME] ("Buyer") hereby submits this non-binding Letter of Intent ("LOI") to acquire the property described above (the "Property"). The Property comprises approximately **{acres:.0f} acres** of land located in {county} County, Texas, within the ERCOT {ercot} risk zone and within {dist} miles of existing 345 kV transmission infrastructure.

**Indicative Purchase Price:** {price_fmt} ({ppa_fmt})

---

## 2. Key Terms

| Term | Detail |
|------|--------|
| **Property Size** | ~{acres:.0f} acres |
| **Land Use** | {land_use} |
| **Existing Improvements** | None (vacant land) |
| **Zoning** | {zoning} |
| **Transmission Access** | {dist} mi from 345 kV line |
| **Zoning/Power Synergy** | +{synergy} pts |
| **ERCOT Risk** | {ercot} |
| **ERCOT Region** | {ercot_zone} |

---

## 3. Texas Congestion Exemption Clause

The Parties acknowledge that the Property is situated within the ERCOT interconnection queue area. Buyer's intended use involves on-site generation and/or behind-the-meter power that does not rely on ERCOT grid connectivity for primary operations.

**Notwithstanding any provision to the contrary, Seller shall have no liability, and Buyer assumes all risk, for:**

(a) ERCOT transmission congestion or curtailment;
(b) Changes to ERCOT interconnection rules or queue priority under PURPA, Section 210, or Texas Utilities Code Section 35.004;
(c) Remote disconnection or load-shed events ordered by ERCOT under NERC Standard EOP-003 or Texas Administrative Code Section 25.361;
(d) Any delay, cost, or loss arising from the above, including but not limited to lost revenue, stranded assets, or forced curtailment of operations.

This clause survives closing and is not subject to the survival limitations in Section 8. Reference is made to Texas Senate Bill 6 (2025) regarding critical infrastructure load protection.

---

## 4. Due Diligence Period

Buyer shall have **90 days** from execution to conduct:
- Phase I & II environmental assessment
- Geotechnical survey
- Title review and survey
- ERCOT interconnection study
- Transmission capacity analysis (onsite 345 kV access verification)
- Floodplain and wetland delineation
- Zoning and entitlement verification

---

## 5. Exclusivity

Seller agrees not to solicit or entertain alternative offers for the Property during the Due Diligence Period.

---

## 6. Confidentiality

This LOI and all related negotiations are confidential. Neither party shall disclose terms without prior written consent, except as required by law.

---

## 7. Non-Binding Effect

This LOI is for discussion purposes only. No binding obligation exists unless and until a definitive Purchase and Sale Agreement is executed by both parties. Only Section 6 (Confidentiality) and this Section 7 shall survive.

---

**Submitted by:**

____________________________
[Buyer Name]
[Title]
[Contact Information]

**Acknowledged:**

____________________________
{owner}
Date: _______________
"""

    safe_name = re.sub(r'[^\w]+', '_', owner.strip())[:60]
    path = LOI_DIR / f"loi_{loi_num:02d}_{safe_name}.docx"
    path.parent.mkdir(parents=True, exist_ok=True)

    md_to_docx(template, path)
    return path, template


def main():
    p = argparse.ArgumentParser()
    p.add_argument('--all', action='store_true', help='Generate LOIs for all CRM targets')
    p.add_argument('--top', type=int, default=20, help='Number of LOIs to generate')
    args = p.parse_args()

    targets = load_crm(CRM_PATH)
    if not args.all:
        targets = targets[:args.top]

    print(f"Generating {len(targets)} LOIs...")

    for i, target in enumerate(targets, 1):
        path, _ = fill_loi_target(target, i)
        owner = target['owner_name'].strip()[:40]
        acres = target.get('total_acres') or target.get('primary_acres', '0')
        print(f"  {i:>2}. {owner:45s} {acres:>7s}ac → {path.name}")

    print(f"\nDone. {len(targets)} LOIs in {LOI_DIR}/")


if __name__ == "__main__":
    main()
